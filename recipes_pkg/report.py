"""Модуль для сохранения отчётов"""

from docx import Document
from openpyxl import Workbook


def save_to_docx(dish_info, filename):
    """Сохраняет отчёт в формате DOCX"""
    doc = Document()
    doc.add_heading(dish_info['name'], 0)
    doc.add_paragraph(f"Стоимость: {dish_info['cost']} руб.")
    n = dish_info['nutrition']
    doc.add_paragraph(f"Калории: {n['calories']} ккал")
    doc.add_paragraph(f"Белки: {n['protein']} г")
    doc.add_paragraph(f"Жиры: {n['fat']} г")
    doc.add_paragraph(f"Углеводы: {n['carbs']} г")
    doc.save(filename)


def save_to_xlsx(dish_info, filename):
    """Сохраняет отчёт в формате XLSX"""
    wb = Workbook()
    ws = wb.active
    ws['A1'] = 'Показатель'
    ws['B1'] = 'Значение'
    ws['A2'] = 'Блюдо'
    ws['B2'] = dish_info['name']
    ws['A3'] = 'Стоимость'
    ws['B3'] = dish_info['cost']
    n = dish_info['nutrition']
    ws['A4'] = 'Калории'
    ws['B4'] = n['calories']
    ws['A5'] = 'Белки'
    ws['B5'] = n['protein']
    ws['A6'] = 'Жиры'
    ws['B6'] = n['fat']
    ws['A7'] = 'Углеводы'
    ws['B7'] = n['carbs']
    wb.save(filename)